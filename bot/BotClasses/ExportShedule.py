import datetime
import os

import docx

from .StudentShedule import StudentShedule
from ics import Calendar, Event

tt_dict = {
    "08:00": "05:00",
    "09:40": "06:40",
    "11:20": "08:20",
    "13:30": "10:30",
    "15:10": "12:10",
    "16:50": "13:50",
    "18:25": "15:25",
    "20:00": "17:00",
    "08:00:00": "05:00:00",
    "09:40:00": "06:40:00",
    "11:20:00": "08:20:00",
    "13:30:00": "10:30:00",
    "15:10:00": "12:10:00",
    "16:50:00": "13:50:00",
    "18:25:00": "15:25:00",
    "20:00:00": "17:00:00",

}

path = ''
try:
    if os.getenv("OS") == 'Windows_NT':  # test key
        path = "templates/"
    else:
        path = '/home/u_botkai/botraspisanie/botkai_telegram/templates/'
except:
    path = '/home/u_botkai/botraspisanie/botkai_telegram/templates/'



class ShedRow(object):
    def __init__(self, daytime, daydate, disciplname, discipltype, auditory, building, prepodName):
        self.daytime = daytime
        self.daydate = daydate
        self.disciplname = disciplname
        self.discipltype = discipltype
        self.auditory = auditory
        self.building = building
        self.prepodName = prepodName


class ExportShedule(StudentShedule):
    async def makeFile(self, week):

        c = Calendar()
        today = datetime.date.today()
        current_date = today - datetime.timedelta(days=today.isoweekday()) + datetime.timedelta(days=1)
        isNormal, response = await super()._get_response()
        if not response:
            return False
        days_in_week = get_unique_daynums(response)
        days_in_week.sort()

        current_week = 0
        while (current_week <= week):
            if str(current_date.isoweekday()) not in days_in_week:
                current_date += datetime.timedelta(days=1)
                continue
            for key in response:
                if (current_date.month == 12 and current_date.day == 30) or (
                        current_date.month == 7 and current_date.day == 1):
                    break
                chetnost = False if (datetime.date(current_date.year, current_date.month,
                                                   current_date.day).isocalendar()[
                                         1] + self.chetn) % 2 else True  # Если True чет, False - неч

                for row in response:
                    daydate = row.get("daydate", "").rstrip().lower()
                    prefix = ""

                    if (daydate == 'чет' and not chetnost) or (daydate == 'неч' and chetnost):
                        continue
                    elif daydate == 'чет/неч':
                        if chetnost:
                            prefix = " (1) гр."
                        else:
                            prefix = " (2) гр."
                    elif daydate == 'неч/чет':
                        if chetnost:
                            prefix = " (2) гр."
                        else:
                            prefix = " (1) гр."

                    e = Event()
                    tt = row.get("daytime", "").rstrip() if len(row.get("daytime", "").rstrip()) < 6 else row.get("daytime", "").rstrip()[:5]
                    tt = tt_dict.get(tt, tt)
                    begin_time = str(current_date) + " {}:00".format(tt)
                    # end_time = str(current_date) + " {}:00".format(time_dict[row.get("daytime", "").rstrip()])
                    e.name = prefix + row.get("discipltype", "").rstrip().upper() + " " + row.get("disciplname", "").rstrip()
                    e.begin = begin_time
                    e.duration = datetime.timedelta(
                        minutes=190 if row.get("discipltype", "").rstrip().upper() == 'Л.Р.' else 90)
                    e.location = "В {} ауд. {} зд".format(row.get("auditory", "").rstrip(), row.get("building", "").rstrip())
                    e.description = "В {} ауд. {} зд".format(row.get("auditory", "").rstrip(), row.get("building", "").rstrip())
                    c.events.add(e)
                current_date = current_date + datetime.timedelta(days=1)
                if str(current_date.isoweekday()) not in days_in_week:
                    current_date = current_date + datetime.timedelta(days=1)
                    continue
            current_week += 1
        return str(c).encode('UTF-8')
        # with open('{}.ics'.format(self.group_id), 'w', encoding="utf-8") as f:
        #     f.write(str(c))

    async def TimetableWrite(self):
        isNormal, response = await super()._get_response()
        if not isNormal:
            return response

        rows = 0
        lis = []

        elem = response[str(1)]
        try:
            lis.append("Понедельник")
            for day in elem:
                lis.append(
                    ShedRow(day["daytime"], day["daydate"], day["disciplname"], day["discipltype"], day["auditory"],
                            day["building"], day["prepodName"]))
            elem = response[str(2)]
            lis.append("Вторник")
            for day in elem:
                lis.append(
                    ShedRow(day["daytime"], day["daydate"], day["disciplname"], day["discipltype"], day["auditory"],
                            day["building"], day["prepodName"]))
            rows = len(lis)
            elem = response[str(3)]
            lis.append("Среда")
            for day in elem:
                lis.append(
                    ShedRow(day["daytime"], day["daydate"], day["disciplname"], day["discipltype"], day["auditory"],
                            day["building"], day["prepodName"]))
            rows = len(lis)
            elem = response[str(4)]
            lis.append("Четверг")
            for day in elem:
                lis.append(
                    ShedRow(day["daytime"], day["daydate"], day["disciplname"], day["discipltype"], day["auditory"],
                            day["building"], day["prepodName"]))
            rows = len(lis)
            elem = response[str(5)]
            lis.append("Пятница")
            for day in elem:
                lis.append(
                    ShedRow(day["daytime"], day["daydate"], day["disciplname"], day["discipltype"], day["auditory"],
                            day["building"], day["prepodName"]))
            rows = len(lis)
            elem = response[str(6)]
            lis.append("Суббота")
            for day in elem:
                lis.append(
                    ShedRow(day["daytime"], day["daydate"], day["disciplname"], day["discipltype"], day["auditory"],
                            day["building"], day["prepodName"]))
        except:
            pass
        return lis

    async def createDocShedule(self):

        wordDocument = docx.Document(path + "blank.docx")

        lis = await self.TimetableWrite()
        for day in lis:
            if day in ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота"]:
                par = wordDocument.add_heading(day, 3)
                par.bold = True


            else:
                par = wordDocument.add_paragraph(
                    (str(day.daytime)).rstrip() + " " + ((str(day.daydate)).rstrip()).ljust(8) + " " + str(
                        day.disciplname) + " " + (str(day.discipltype)).upper() + " " + (
                        str(day.auditory)).rstrip() + " ауд  " + (str(day.building)).rstrip() + "зд.  " + (
                        str(day.prepodName)).rstrip())
                # par.style = "No Spacing"
        wordDocument.save(path + str(self.group_id) + ".docx")
        return path + str(self.group_id) + ".docx"


def get_unique_daynums(objects):
    unique_daynums = set()
    for obj in objects:
        daynum = obj.get("daynum")
        if daynum:
            unique_daynums.add(daynum)
    return list(unique_daynums)